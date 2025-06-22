"""Extractor for Microsoft Word .docx files."""

from __future__ import annotations

from pathlib import Path
import typing as _t

import docx  # type: ignore  # noqa: F401

from .base_extractor import BaseExtractor, DocumentType

__all__ = ["DocxExtractor"]


class DocxExtractor(BaseExtractor):
    DOCUMENT_TYPE = DocumentType.DOCX

    def can_process(self, source: _t.Union[str, Path]) -> bool:
        return str(source).lower().endswith(".docx")

    def extract_text(self, source: _t.Union[str, Path, bytes]) -> str:  # noqa: D401
        """Return text from a .docx file.

        The routine iterates over *all* paragraphs *and* table cells, preserving
        logical reading order.  Paragraphs are separated by newlines; table
        cells are joined by tabs within a row, and rows by newlines.
        """

        from io import BytesIO  # local import to keep global namespace minimal

        # Obtain python-docx `Document` instance from various input shapes
        try:
            if isinstance(source, (bytes, bytearray)):
                doc_obj = docx.Document(BytesIO(source))  # type: ignore[arg-type]
            else:
                path = self._to_path(source)
                doc_obj = docx.Document(str(path))
        except Exception as exc:  # pragma: no cover – propagate meaningful message
            raise RuntimeError("Failed to open DOCX document") from exc

        # Collect text from paragraphs
        parts: list[str] = []

        for para in doc_obj.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Collect text from tables (row-major order)
        for table in doc_obj.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts) 